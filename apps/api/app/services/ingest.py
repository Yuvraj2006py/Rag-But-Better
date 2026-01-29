from __future__ import annotations

import os
from pathlib import Path

from bs4 import BeautifulSoup
from pypdf import PdfReader
from docx import Document as DocxDocument

from app.core.models import Document


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _read_html(path: Path) -> str:
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text("\n")


def parse_file(path: str) -> Document:
    file_path = Path(path)
    ext = file_path.suffix.lower()
    if ext in {".txt", ".md"}:
        text = _read_txt(file_path)
    elif ext == ".pdf":
        text = _read_pdf(file_path)
    elif ext == ".docx":
        text = _read_docx(file_path)
    elif ext in {".html", ".htm"}:
        text = _read_html(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return Document(
        doc_id=file_path.stem,
        title=file_path.name,
        source_path=os.path.abspath(path),
        text=text,
        metadata={"ext": ext},
    )
